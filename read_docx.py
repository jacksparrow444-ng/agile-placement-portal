from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json

doc = Document('Sprint_3_Testing_Report.docx')

output = []

# Core document properties
output.append(f"=== DOCUMENT SECTIONS: {len(doc.sections)} ===")
sec = doc.sections[0]
output.append(f"Page size: {sec.page_width.inches:.2f}\" x {sec.page_height.inches:.2f}\"")
output.append(f"Margins: top={sec.top_margin.inches:.2f}\" bottom={sec.bottom_margin.inches:.2f}\" left={sec.left_margin.inches:.2f}\" right={sec.right_margin.inches:.2f}\"")

output.append(f"\n=== PARAGRAPHS ({len(doc.paragraphs)}) ===")
for i, para in enumerate(doc.paragraphs[:60]):
    style_name = para.style.name
    text = para.text[:100]
    align = para.alignment
    
    # Get first run properties
    run_info = ""
    if para.runs:
        r = para.runs[0]
        bold = r.bold
        size = r.font.size.pt if r.font.size else "inherit"
        color = str(r.font.color.rgb) if r.font.color and r.font.color.type else "inherit"
        run_info = f" | bold={bold} size={size} color={color}"
    
    output.append(f"P{i:03d} [{style_name}] align={align}: {repr(text)}{run_info}")

output.append(f"\n=== TABLES ({len(doc.tables)}) ===")
for ti, table in enumerate(doc.tables):
    output.append(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    # Print first 3 rows of each table
    for ri, row in enumerate(table.rows[:4]):
        row_data = []
        for cell in row.cells:
            # Get cell background
            cell_text = cell.text[:30].strip()
            row_data.append(repr(cell_text))
        output.append(f"  Row {ri}: {row_data}")
    if len(table.rows) > 4:
        output.append(f"  ... ({len(table.rows)-4} more rows)")

# Write to file
with open('sprint3_structure.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(output))
print("DONE")
